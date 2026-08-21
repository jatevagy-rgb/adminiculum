#!/usr/bin/env python3
"""Generate the non-production, source-grounded first compliance template batch."""
from datetime import date
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "artifacts" / "compliance-templates"
MANIFEST_VERSION = "1.1.0"
CORPUS_MANIFEST_VERSION = "1.1.0"
STATUS = "LEGAL_REVIEW_REQUIRED"
FIELDS = [
    "company.legal_name", "company.registered_office", "company.registration_number",
    "company.tax_number", "company.employee_count", "company.main_activity",
    "company.representative_name", "document.version", "document.effective_date",
    "document.review_date", "document.owner_name", "document.approved_by"
]

TEMPLATES = [
    {
        "key": "gdpr-processing-activities-register", "title": "Adatkezelési tevékenységek nyilvántartása",
        "subtitle": "Adatkezelői nyilvántartás - felülvizsgálati tervezet", "kind": "register",
        "requirements": ["GDPR_ROPA_CONTROLLER"], "anchors": ["A-GDPR-30-1-5"],
        "sections": [
            ("Rendeltetés", "A nyilvántartás az adatkezelő felelősségébe tartozó adatkezelési tevékenységek áttekinthető, írásbeli rögzítésére szolgál."),
            ("Nyilvántartási bejegyzés", "Minden adatkezelési tevékenységet külön, a tényleges működésnek megfelelő bejegyzésben kell rögzíteni."),
            ("Felülvizsgálat", "A dokumentumgazda a változásokat és a jogi felülvizsgálat megállapításait a nyilvántartásban vezeti át.")
        ],
        "table": ["Adatkezelési cél", "Érintettek és adatkategóriák", "Címzettek / továbbítás", "Megőrzés", "Biztonsági intézkedések"],
        "extra_fields": ["privacy.processing_purpose", "privacy.legal_basis"]
    },
    {
        "key": "gdpr-direct-collection-privacy-notice", "title": "Adatkezelési tájékoztató",
        "subtitle": "Közvetlen adatgyűjtéshez - jogi felülvizsgálati tervezet", "kind": "notice",
        "requirements": ["GDPR_PRIVACY_INFORMATION_DIRECT_COLLECTION"], "anchors": ["A-GDPR-13-1-2"],
        "sections": [
            ("Ki kezeli az adatokat", "Adatkezelő: {{company.legal_name}}; székhely: {{company.registered_office}}."),
            ("Az adatkezelés célja és jogalapja", "Cél: {{privacy.processing_purpose}}. Jogalap: {{privacy.legal_basis}}."),
            ("Érintetti jogok és jogorvoslat", "Az érintett a rá vonatkozó jogokat az adatkezelőnél gyakorolhatja. A konkrét tájékoztató tartalmát a tényleges adatkezeléshez kell igazítani."),
            ("Kapcsolat", "A kérelmek fogadására kijelölt kapcsolattartót és elérhetőséget a végleges, jóváhagyott változatban kell rögzíteni.")
        ], "table": None, "extra_fields": ["privacy.processing_purpose", "privacy.legal_basis"]
    },
    {
        "key": "consumer-complaint-record", "title": "Fogyasztói panasz jegyzőkönyve",
        "subtitle": "Kitöltendő nyilvántartási űrlap", "kind": "form",
        "requirements": ["CONSUMER_COMPLAINT_RECORD_AND_RESPONSE"], "anchors": ["A-CP-17A-3A-8"],
        "sections": [("Használat", "A jegyzőkönyvet a fogyasztóvédelmi szabályok szerinti esetekben, a tényleges panaszkezelési folyamathoz igazítva kell kitölteni.")],
        "table": ["Fogyasztó neve és elérhetősége", "A panasz előterjesztésének helye, ideje és módja", "A panasz részletes leírása és bizonyítékai", "A vállalkozás álláspontja", "Jegyzőkönyv felvevője", "Egyedi azonosító", "Fogyasztó aláírása / nyilatkozata"],
        "extra_fields": ["complaint.reference_number"]
    },
    {
        "key": "consumer-complaint-written-response", "title": "Írásbeli válasz fogyasztói panaszra",
        "subtitle": "Ügyfélnek küldhető választervezet - jogi felülvizsgálat szükséges", "kind": "letter",
        "requirements": ["CONSUMER_COMPLAINT_RECORD_AND_RESPONSE"], "anchors": ["A-CP-17A-3A-8"],
        "sections": [
            ("Hivatkozás", "Ügyszám: {{complaint.reference_number}}. Tisztelt Fogyasztó!"),
            ("A panasz kivizsgálása", "A bejelentésben foglaltakat megvizsgáltuk. A végleges válaszban a tényállást, a döntést és annak indokait pontosan kell rögzíteni."),
            ("Döntés és intézkedés", "[JOGÁSZI KITÖLTÉS: elfogadás / részbeni elfogadás / elutasítás; intézkedés és határidő]"),
            ("Jogorvoslati tájékoztatás", "Elutasítás esetén a végleges levélben a vonatkozó hatósági vagy békéltető testületi tájékoztatást a konkrét ügyhöz igazítva kell megadni.")
        ], "table": None, "extra_fields": ["complaint.reference_number"]
    },
    {
        "key": "whistleblowing-internal-procedure", "title": "Belső visszaélés-bejelentési eljárásrend",
        "subtitle": "Foglalkoztatói működési tervezet - jogi felülvizsgálat szükséges", "kind": "policy",
        "requirements": ["WHISTLEBLOWING_INTERNAL_CHANNEL", "WHISTLEBLOWING_ACKNOWLEDGEMENT_AND_INVESTIGATION"], "anchors": ["A-WB-18-19", "A-WB-21-25"],
        "sections": [
            ("Cél és hatály", "Az eljárásrend a {{company.legal_name}} belső visszaélés-bejelentési rendszerének működési keretét rögzíti."),
            ("Működtető és pártatlanság", "A rendszer működtetőjét, helyettesítését és összeférhetetlenségi szabályait a jóváhagyott változatban kell kijelölni."),
            ("Bejelentés módja", "A bejelentés írásban vagy szóban tehető. A hozzáférési csatornákat a vállalkozás tényleges működéséhez kell igazítani."),
            ("Visszaigazolás és kivizsgálás", "Írásbeli bejelzésnél a visszaigazolást hét napon belül kell elküldeni. A kivizsgálás főszabály szerint legfeljebb harminc nap; indokolt hosszabbítás esetén a törvényi garanciákat kell alkalmazni."),
            ("Bizalmasság és iratkezelés", "A személyazonosságot és a bejelentés tartalmát a jogosultakon kívül más nem ismerheti meg; az iratkezelés részletszabályait külön felül kell vizsgálni."),
            ("Tájékoztatás és felülvizsgálat", "A rendszer működéséről világos és könnyen hozzáférhető tájékoztatást kell nyújtani. A dokumentumot változás esetén felül kell vizsgálni.")
        ], "table": None, "extra_fields": []
    },
    {
        "key": "whistleblowing-written-acknowledgement", "title": "Visszaigazolás belső visszaélés-bejelentésről",
        "subtitle": "Írásbeli visszaigazolási tervezet", "kind": "letter",
        "requirements": ["WHISTLEBLOWING_ACKNOWLEDGEMENT_AND_INVESTIGATION"], "anchors": ["A-WB-21-25"],
        "sections": [
            ("Hivatkozás", "Bejelentés azonosítója: {{whistleblowing.report_reference}}. Tisztelt Bejelentő!"),
            ("Visszaigazolás", "Ezúton visszaigazoljuk bejelentésének kézhezvételét."),
            ("Eljárási tájékoztatás", "A bejelentést a vonatkozó belső eljárásrend és a jogszabályi keretek szerint vizsgáljuk ki. A személyes adatok kezelésére vonatkozó tájékoztatást a végleges dokumentumhoz kell csatolni vagy elérhetővé tenni."),
            ("Kapcsolattartás", "A vizsgálat során szükség esetén további információt kérhetünk. A megkeresésekre a kijelölt csatornán válaszoljon.")
        ], "table": None, "extra_fields": ["whistleblowing.report_reference"]
    }
]

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); r = p.add_run(text); r.bold = bold; r.font.name = 'Arial'; r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial'); r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def style_doc(doc):
    sec = doc.sections[0]; sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.7); sec.left_margin=Cm(2.1); sec.right_margin=Cm(2.1)
    normal=doc.styles['Normal']; normal.font.name='Arial'; normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for name,size,color in [('Heading 1',16,'17365D'),('Heading 2',12,'17365D')]:
        st=doc.styles[name]; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('ADMINICULUM | JOGI FELÜLVIZSGÁLATI TERVEZET'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('LEGAL_REVIEW_REQUIRED | Belső használatra | '); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

def metadata(doc):
    table=doc.add_table(rows=3, cols=2); table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.style='Table Grid'
    values=[('Szervezet','{{company.legal_name}}'),('Verzió','{{document.version}}'),('Hatálybalépés','{{document.effective_date}}'),('Dokumentumgazda','{{document.owner_name}}'),('Jóváhagyó','{{document.approved_by}}'),('Felülvizsgálat','{{document.review_date}}')]
    for i,(label,value) in enumerate(values):
        row=i//2; col=(i%2)*1; cell=table.cell(row,col); set_cell_text(cell,label+'\n'+value); shade(cell,'F2F5F8')
    doc.add_paragraph()

def add_title(doc, template):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run(template['title']); r.font.name='Arial'; r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=RGBColor(23,54,93)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14); r=p.add_run(template['subtitle']); r.font.name='Arial'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(89,89,89)

def add_table(doc, labels):
    table=doc.add_table(rows=1, cols=2); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_cell_text(table.rows[0].cells[0], 'Kötelező tartalmi elem', True); set_cell_text(table.rows[0].cells[1], 'Rögzítendő adat / megjegyzés', True)
    for c in table.rows[0].cells: shade(c,'DDE7F1')
    for label in labels:
        cells=table.add_row().cells; set_cell_text(cells[0], label, True); set_cell_text(cells[1], '[KÉZI KITÖLTÉS SZÜKSÉGES]')
    doc.add_paragraph()

def make_doc(template, target):
    doc=Document(); style_doc(doc); add_title(doc,template); metadata(doc)
    if template['table']: add_table(doc,template['table'])
    for idx,(heading,body) in enumerate(template['sections'],1):
        doc.add_heading(f'{idx}. {heading}', level=1); doc.add_paragraph(body)
    doc.add_heading('Jóváhagyás',level=1); doc.add_paragraph('Jóváhagyó: {{document.approved_by}}\nDátum: {{document.effective_date}}\nAláírás: ______________________________')
    doc.save(target)

def package(template, reqs):
    folder=OUT/template['key']; folder.mkdir(parents=True,exist_ok=True)
    make_doc(template,folder/'template.docx')
    used=FIELDS+template['extra_fields']
    spec={"templateKey":template['key'],"version":"0.1.0","generatedAt":"2026-08-21","sourceCorpusManifestVersion":CORPUS_MANIFEST_VERSION,"legalReviewStatus":STATUS,"approvedAt":None,"supersedes":None,"legalTextPolicy":"FIXED_TEXT_REVIEW_REQUIRED","companySpecificFields":used,"lawyerSelectableOptions":[],"conditionalSections":[{"sectionKey":"applicability.review","conditionRuleKey":"MANUAL_LEGAL_REVIEW","variant":"REQUIRED","legalReviewStatus":STATUS}],"rendering":"Deterministic replacement only; no runtime AI; output is a reviewable draft."}
    basis={"templateKey":template['key'],"legalReviewStatus":STATUS,"requirementKeys":template['requirements'],"sourceAnchors":[reqs['anchors'][a] | {"anchorKey":a} for a in template['anchors']],"customerDocumentCitation":"Human-readable legal references may be added only after lawyer review; hashes remain internal."}
    (folder/'template-spec.json').write_text(json.dumps(spec,ensure_ascii=False,indent=2),encoding='utf-8')
    (folder/'legal-basis.json').write_text(json.dumps(basis,ensure_ascii=False,indent=2),encoding='utf-8')
    (folder/'fields.json').write_text(json.dumps({"templateKey":template['key'],"registryFile":"../template-fields.json","fieldKeys":used},ensure_ascii=False,indent=2),encoding='utf-8')
    (folder/'REVIEW.md').write_text(f"# {template['title']}\n\nStátusz: `{STATUS}`.\n\nA sablon csak forrás-ankorolt candidate követelményekhez készült. Jóváhagyás előtt jogásznak kell ellenőriznie a tényleges alkalmazhatóságot, a kitöltött tényállást és a végleges szöveget.\n",encoding='utf-8')

def main():
    reqs=json.loads((ROOT/'docs/compliance/legal-review/requirements-candidates.json').read_text(encoding='utf-8'))
    for template in TEMPLATES: package(template,reqs)
    (OUT/'README.md').write_text('# Compliance template candidates\n\nNon-production, source-grounded candidate packages. Every template is `LEGAL_REVIEW_REQUIRED`; no template is approved or automatically renderable for customer publication.\n',encoding='utf-8')

if __name__=='__main__': main()
