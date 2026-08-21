#!/usr/bin/env python3
"""Deterministic non-production DOCX token renderer proof.

Only whole text-node tokens are replaced. Unknown or missing required tokens fail.
The source template is opened read-only and the result is always a new file.
"""
import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from docx import Document

TOKEN = re.compile(r"\{\{([a-z0-9_.]+)\}\}")

def tokens_in_document(doc):
    found = []
    for paragraph in list(doc.paragraphs) + [p for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]:
        for run in paragraph.runs:
            found.extend(TOKEN.findall(run.text))
    return found

def render(source, output, fields_path, values_path):
    registry = {field["key"]: field for field in json.loads(Path(fields_path).read_text(encoding="utf-8"))["fields"]}
    values = json.loads(Path(values_path).read_text(encoding="utf-8"))
    doc = Document(source)
    tokens = tokens_in_document(doc)
    unknown = sorted(set(token for token in tokens if token not in registry))
    if unknown:
        raise ValueError("Unknown template fields: " + ", ".join(unknown))
    missing = sorted(set(token for token in tokens if registry[token]["required"] and not values.get(token)))
    if missing:
        raise ValueError("Missing required values: " + ", ".join(missing))
    for paragraph in list(doc.paragraphs) + [p for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]:
        for run in paragraph.runs:
            if TOKEN.search(run.text):
                run.text = TOKEN.sub(lambda match: str(values.get(match.group(1), "[KÉZI KITÖLTÉS SZÜKSÉGES]")), run.text)
    remaining = tokens_in_document(doc)
    if remaining:
        raise ValueError("Partial token replacement detected: " + ", ".join(sorted(set(remaining))))
    Path(output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("source")
    parser.add_argument("output")
    parser.add_argument("--fields", required=True)
    parser.add_argument("--values", required=True)
    args = parser.parse_args()
    render(args.source, args.output, args.fields, args.values)
